from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

INPUT_MD = Path('HappyPaws_V3_2_Definitiva_Corregida.md')
OUTPUT_DOCX = Path('HappyPaws_V3_2_Definitiva_Corregida.docx')


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def format_text_run(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Arial'
        run.font.size = Pt(11)


def add_text_paragraph(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    format_text_run(p, text)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.style = 'Title'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.style = f'Heading {level}'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(18 if level == 0 else 14 if level == 1 else 12)
    return p


def parse_table(lines, start_index):
    rows = []
    i = start_index
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.fullmatch(r'\|(?:\s*[:-]?[-]+[:-]?\s*\|)+', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    headers = rows[0]
    data = rows[1:] if len(rows) > 1 else []
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, 'D9EAF7')
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(10)
    for row in data:
        cells = table.add_row().cells
        for idx in range(len(headers)):
            value = row[idx] if idx < len(row) else ''
            cells[idx].text = value
            set_cell_margins(cells[idx])
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
    return table


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = INPUT_MD.read_text(encoding='utf-8').splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 0)
            i += 1
            continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith('|'):
            table_rows, i = parse_table(lines, i)
            add_table(doc, table_rows)
            continue
        if line.startswith('**') and line.endswith('**') and len(line) > 4:
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            i += 1
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            format_text_run(p, line[2:].strip())
            i += 1
            continue
        add_text_paragraph(doc, line)
        i += 1

    doc.save(OUTPUT_DOCX)
    print(f'Generado: {OUTPUT_DOCX}')


if __name__ == '__main__':
    build_doc()
