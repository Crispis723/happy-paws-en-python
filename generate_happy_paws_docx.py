from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_FILE = 'HappyPaws_V3_1_Definitiva_Corregida.docx'
UPDATE_DATE = '30 de mayo de 2026'
PROJECT_START = '15 de junio de 2026'
FINAL_DELIVERY = '16 de agosto de 2026'
BASE_COST = 16400000
CONTINGENCY = 2460000
TOTAL_COST = 18860000
SALE_PRICE = 24518000
ORACLE_EXTRA = 1200000
TOTAL_WITH_ORACLE = 20060000
SALE_WITH_ORACLE = 26078000


def format_cop(value):
    return f'COP {value:,.0f}'.replace(',', '.')


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = f'Happy Paws | Versión 3.1 | Actualizado: {UPDATE_DATE}'
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)


def set_document_style(doc):
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(11)
    for name in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[name].font.name = 'Arial'
        styles[name].font.bold = True
        styles[name].font.size = Pt(14)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        add_header(section)


def add_title(doc, text, size=18):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f'Heading {level}'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.italic = True if level >= 2 else False
    run.font.name = 'Arial'
    run.font.size = Pt(14 if level == 1 else 12)
    return paragraph


def add_paragraph(doc, text=''):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style='List Bullet')
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style='List Number')
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return paragraph


def add_label_value(doc, label, value):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f'{label}: ')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run2 = paragraph.add_run(value)
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def add_pipe_block(doc, headers, rows):
    separator = '| ' + ' | '.join(['---'] * len(headers)) + ' |'
    add_paragraph(doc, '| ' + ' | '.join(headers) + ' |')
    add_paragraph(doc, separator)
    for row in rows:
        add_paragraph(doc, '| ' + ' | '.join(str(item) for item in row) + ' |')


def add_note(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f'Nota: {text}')
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    return paragraph


def section_title(doc, number, title):
    add_heading(doc, f'{number}. {title.upper()}', 1)


def subsection_title(doc, number, title):
    add_heading(doc, f'{number} {title}', 2)


doc = Document()
set_document_style(doc)

# Portada
add_title(doc, 'HAPPY PAWS')
add_title(doc, 'VERSIÓN 3.1 - DEFINITIVA CORREGIDA', 16)
add_paragraph(doc, 'Documento formal, ejecutivo, legal y ejecutable listo para copiar a Word.')
for line in [
    f'Fecha de actualización: {UPDATE_DATE}',
    'Cliente: Veterinaria V+Kotiando',
    'Proyecto: Plataforma SaaS multiempresa para bienestar animal en Colombia',
    'Stack: React, Spring Boot, PostgreSQL, JWT + Spring Security, Docker, Cloudinary, AWS/Railway'
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line)
doc.add_page_break()

# Índice
section_title(doc, 1, 'Índice general')
for item in [
    'Resumen ejecutivo',
    'Módulos obligatorios del MVP',
    'Stack tecnológico y arquitectura',
    'Costos corregidos del proyecto',
    'Cronograma con fechas reales',
    'Aspectos legales colombianos',
    'Seguridad adicional y cumplimiento técnico',
    'Story points realistas',
    'Endpoints REST y códigos de error',
    'Análisis de competencia',
    'Análisis de riesgos legales y técnicos',
    'Definición de Done mejorada',
    'Qué no incluir en el MVP',
    'Conclusión y nota final',
    'Anexos'
]:
    add_numbered(doc, item)
add_note(doc, 'En Word, este índice puede actualizarse con estilos de título si se desea tabla de contenido automática.')

# Resumen ejecutivo
section_title(doc, 2, 'Resumen ejecutivo')
add_paragraph(doc, 'Happy Paws es una plataforma SaaS multiempresa para bienestar animal en Colombia. Conecta dueños de mascotas, veterinarias, refugios y entidades de bienestar animal en un solo ecosistema digital orientado a citas, historiales, adopciones, reportes ciudadanos y control administrativo.')
add_label_value(doc, 'Misión', 'Digitalizar la gestión veterinaria y la protección animal con una solución confiable, trazable y accesible para Colombia.')
add_label_value(doc, 'Visión', 'Ser la plataforma de referencia en Latinoamérica para la gestión integral del bienestar animal.')
add_label_value(doc, 'Propuesta de valor', 'Una sola plataforma para gestionar mascotas, citas, adopciones, reportes y operación veterinaria con enfoque local.')
add_table(doc, ['Elemento', 'Valor / alcance', 'Observación'], [
    ['Mercado objetivo inicial', 'Bogotá, Medellín e Ibagué', 'Entrada comercial por ciudades con alta densidad de mascotas'],
    ['Enfoque de negocio', 'SaaS B2B + componente social B2C', 'Vets y refugios como clientes principales'],
    ['Inversión requerida', format_cop(TOTAL_COST), 'Costo total MVP corregido'],
    ['Precio sugerido al cliente', format_cop(SALE_PRICE), 'Margen comercial de 30%']
])
add_pipe_block(doc, ['Elemento', 'Valor / alcance', 'Observación'], [
    ['Mercado objetivo inicial', 'Bogotá, Medellín e Ibagué', 'Entrada comercial por ciudades con alta densidad de mascotas'],
    ['Enfoque de negocio', 'SaaS B2B + componente social B2C', 'Vets y refugios como clientes principales'],
    ['Inversión requerida', format_cop(TOTAL_COST), 'Costo total MVP corregido'],
    ['Precio sugerido al cliente', format_cop(SALE_PRICE), 'Margen comercial de 30%']
])

# MVP
section_title(doc, 3, 'Módulos obligatorios del MVP')
add_table(doc, ['Módulo', 'Incluye', 'Resultado esperado'], [
    ['Registro e inicio de sesión', 'JWT + roles: Administrador, Veterinaria, Propietario, Refugio', 'Acceso seguro y segmentado'],
    ['CRUD de mascotas', 'Nombre, raza, edad, peso, foto, vacunas', 'Ficha unificada por mascota'],
    ['Agenda de citas', 'Crear, cancelar, reprogramar, notificar', 'Control operativo de agenda'],
    ['Historial clínico', 'Consultas, diagnósticos, medicamentos, vacunas', 'Trazabilidad de atención médica'],
    ['Módulo de adopciones', 'Publicar, postular, aprobar, rechazar', 'Flujo digital de adopción'],
    ['Reportes ciudadanos', 'Maltrato o abandono con geolocalización y seguimiento', 'Canal de protección animal'],
    ['Panel administrativo', 'Métricas, usuarios, moderación', 'Control del ecosistema']
])
add_pipe_block(doc, ['Módulo', 'Incluye', 'Resultado esperado'], [
    ['Registro e inicio de sesión', 'JWT + roles: Administrador, Veterinaria, Propietario, Refugio', 'Acceso seguro y segmentado'],
    ['CRUD de mascotas', 'Nombre, raza, edad, peso, foto, vacunas', 'Ficha unificada por mascota'],
    ['Agenda de citas', 'Crear, cancelar, reprogramar, notificar', 'Control operativo de agenda'],
    ['Historial clínico', 'Consultas, diagnósticos, medicamentos, vacunas', 'Trazabilidad de atención médica'],
    ['Módulo de adopciones', 'Publicar, postular, aprobar, rechazar', 'Flujo digital de adopción'],
    ['Reportes ciudadanos', 'Maltrato o abandono con geolocalización y seguimiento', 'Canal de protección animal'],
    ['Panel administrativo', 'Métricas, usuarios, moderación', 'Control del ecosistema']
])

# Stack
section_title(doc, 4, 'Stack tecnológico y arquitectura')
add_paragraph(doc, 'La arquitectura recomendada es modular, segura y escalable. Se prioriza desacoplar frontend, backend y almacenamiento de imágenes para facilitar mantenimiento y despliegue.')
add_table(doc, ['Capa', 'Tecnología', 'Uso'], [
    ['Frontend', 'React', 'Interfaz web responsiva y mantenible'],
    ['Backend', 'Spring Boot', 'API REST principal del negocio'],
    ['Base de datos', 'PostgreSQL (base) / Oracle Database (opcional empresarial)', 'Persistencia relacional y alternativa corporativa'],
    ['Autenticación', 'JWT + Spring Security', 'Seguridad y control de acceso'],
    ['Contenedores', 'Docker', 'Portabilidad y despliegue reproducible'],
    ['Imágenes', 'Cloudinary', 'Gestión de fotos de mascotas y archivos'],
    ['Despliegue', 'AWS / Railway', 'Publicación y escalamiento inicial']
])
add_pipe_block(doc, ['Capa', 'Tecnología', 'Uso'], [
    ['Frontend', 'React', 'Interfaz web responsiva y mantenible'],
    ['Backend', 'Spring Boot', 'API REST principal del negocio'],
    ['Base de datos', 'PostgreSQL (base) / Oracle Database (opcional empresarial)', 'Persistencia relacional y alternativa corporativa'],
    ['Autenticación', 'JWT + Spring Security', 'Seguridad y control de acceso'],
    ['Contenedores', 'Docker', 'Portabilidad y despliegue reproducible'],
    ['Imágenes', 'Cloudinary', 'Gestión de fotos de mascotas y archivos'],
    ['Despliegue', 'AWS / Railway', 'Publicación y escalamiento inicial']
])
add_table(doc, ['Componente', 'Decisión técnica', 'Motivo'], [
    ['API', 'REST versionada /api/v1', 'Facilita evolución sin romper integraciones'],
    ['Seguridad', 'Cifrado de contraseñas con bcrypt', 'Protección robusta de credenciales'],
    ['Datos', 'Validación backend obligatoria', 'No depender solo del frontend'],
    ['Auditoría', 'Logs de acciones críticas', 'Trazabilidad legal y operativa']
])
add_table(doc, ['Gestor de base de datos', 'Costo estimado', 'Observación'], [
    ['PostgreSQL', format_cop(0), 'Opción recomendada por costo y escalabilidad para el MVP'],
    ['Oracle Database administrado', format_cop(ORACLE_EXTRA), 'Alternativa empresarial estimada basada en mercado colombiano 2025-2026']
])

# Costos
section_title(doc, 5, 'Costos corregidos del proyecto')
add_paragraph(doc, 'Los costos originales estaban subestimados porque no reflejaban tarifas reales de mercado para Colombia 2025-2026, ni incluían infraestructura mensual, contingencia suficiente, ni el costo real de perfiles junior-mid con experiencia en React, Spring Boot, QA, diseño UX/UI y DevOps.')
add_table(doc, ['Rol / rubro', 'Horas', 'Tarifa por hora', 'Subtotal'], [
    ['Scrum Master', '60', format_cop(35000), format_cop(2100000)],
    ['Desarrollador Frontend (React)', '100', format_cop(40000), format_cop(4000000)],
    ['Desarrollador Backend (Spring Boot)', '100', format_cop(40000), format_cop(4000000)],
    ['QA / Tester', '50', format_cop(35000), format_cop(1750000)],
    ['Diseñador UX/UI', '70', format_cop(35000), format_cop(2450000)],
    ['DevOps (Docker + despliegue)', '40', format_cop(45000), format_cop(1800000)],
    ['Total recursos humanos', '420', '-', format_cop(BASE_COST)]
])
add_table(doc, ['Infraestructura primer mes', 'Valor'], [
    ['Servidor cloud (Railway/AWS)', format_cop(120000)],
    ['Base de datos administrada', format_cop(80000)],
    ['Cloudinary (imágenes)', format_cop(40000)],
    ['Correo transaccional', format_cop(30000)],
    ['Dominio .com (prorrateado)', format_cop(10000)],
    ['Monitoreo', format_cop(20000)],
    ['Total infraestructura mensual', format_cop(300000)]
])
add_table(doc, ['Concepto', 'Valor'], [
    ['Total recursos humanos', format_cop(BASE_COST)],
    ['Infraestructura primer mes', format_cop(300000)],
    ['Contingencia 15%', format_cop(CONTINGENCY)],
    ['Costo total desarrollo MVP', format_cop(TOTAL_COST)],
    ['Precio sugerido al cliente (margen 30%)', format_cop(SALE_PRICE)]
])
add_paragraph(doc, '¿Por qué estaban subestimados los costos originales? Porque se calculaban con supuestos demasiado bajos para perfiles TI en Colombia, no contemplaban infraestructura mensual realista, omitían contingencia suficiente y no reflejaban el costo de cumplir seguridad, despliegue y soporte básico en producción.')
add_paragraph(doc, 'Si el cliente exige Oracle como gestor de base de datos, debe añadirse una bolsa adicional estimada de COP 1.200.000 para el primer mes del MVP, considerando el servicio administrado, soporte y operación básica.')
add_table(doc, ['Ajuste por Oracle', 'Valor'], [
    ['Costo adicional estimado Oracle Database administrado', format_cop(ORACLE_EXTRA)],
    ['Nuevo costo total MVP con Oracle', format_cop(TOTAL_WITH_ORACLE)],
    ['Nuevo precio sugerido al cliente con Oracle (margen 30%)', format_cop(SALE_WITH_ORACLE)]
])
add_note(doc, 'Se usan valores estimados basados en mercado colombiano 2025-2026.')

# Cronograma
section_title(doc, 6, 'Cronograma con fechas reales')
add_paragraph(doc, f'Fecha de inicio del proyecto: {PROJECT_START}. Fecha de entrega final: {FINAL_DELIVERY}. El cronograma se construye con base en la fecha actual y un arranque a una semana, manteniendo cuatro Sprints de dos semanas con cierre y estabilización posterior.')
add_table(doc, ['Sprint', 'Fechas', 'Objetivo', 'Story Points'], [
    ['Sprint 1', '15/06/2026 - 28/06/2026', 'Base y autenticación', '25 SP'],
    ['Sprint 2', '29/06/2026 - 12/07/2026', 'Mascotas y citas', '28 SP'],
    ['Sprint 3', '13/07/2026 - 26/07/2026', 'Adopciones y reportes', '30 SP'],
    ['Sprint 4', '27/07/2026 - 09/08/2026', 'Panel admin y despliegue', '25 SP']
])
add_pipe_block(doc, ['Sprint', 'Fechas', 'Objetivo', 'Story Points'], [
    ['Sprint 1', '15/06/2026 - 28/06/2026', 'Base y autenticación', '25 SP'],
    ['Sprint 2', '29/06/2026 - 12/07/2026', 'Mascotas y citas', '28 SP'],
    ['Sprint 3', '13/07/2026 - 26/07/2026', 'Adopciones y reportes', '30 SP'],
    ['Sprint 4', '27/07/2026 - 09/08/2026', 'Panel admin y despliegue', '25 SP']
])
add_paragraph(doc, 'La entrega del 16 de agosto de 2026 se reserva como ventana de estabilización, ajustes finales, documentación y cierre contractual.')

# Legal
section_title(doc, 7, 'Aspectos legales colombianos')
add_paragraph(doc, 'Esta sección corrige una omisión crítica del documento anterior: la plataforma debe operar con cumplimiento legal colombiano desde el diseño, no como ajuste posterior.')
subsection_title(doc, '7.1', 'Ley 1581 de 2012 - Protección de Datos Personales')
add_paragraph(doc, 'La Ley 1581 de 2012 regula el tratamiento de datos personales en Colombia. Su principio central es el Habeas Data, entendido como el derecho de toda persona a conocer, actualizar, rectificar y suprimir su información cuando sea procedente.')
add_table(doc, ['Actor', 'Responsabilidad'], [
    ['Responsable del tratamiento', 'Define finalidades, recopila autorizaciones, responde solicitudes y protege la base de datos'],
    ['Encargado del tratamiento', 'Procesa datos por cuenta del responsable, implementa medidas de seguridad y atiende instrucciones técnicas'],
    ['Titular', 'Puede consultar, actualizar, rectificar y solicitar eliminación cuando aplique']
])
add_table(doc, ['Derecho del titular', 'Descripción', 'Plazo de respuesta'], [
    ['Consultar', 'Conocer qué datos se tratan y cómo se usan', '10 días hábiles'],
    ['Actualizar / rectificar', 'Corregir información incompleta o inexacta', '10 días hábiles'],
    ['Eliminar / suprimir', 'Solicitar supresión cuando proceda legalmente', '10 días hábiles']
])
add_note(doc, 'Si la solicitud no puede resolverse en 10 días hábiles, la organización debe responder dentro del plazo legal aplicable y dejar trazabilidad del caso.')
subsection_title(doc, '7.2', 'Aviso de privacidad obligatorio')
add_paragraph(doc, 'AVISO DE PRIVACIDAD. Responsable del tratamiento: Happy Paws SAS o la persona jurídica que opere la plataforma para la veterinaria aliada. Finalidades: gestionar usuarios, mascotas, citas, adopciones, reportes, comunicaciones operativas, soporte, seguridad, analítica y cumplimiento legal. Datos recolectados: nombre, cédula o documento, correo, teléfono, dirección, datos de mascotas, historial clínico y registros asociados al uso del sistema. Derechos: consultar, actualizar, rectificar, suprimir, revocar la autorización y presentar quejas ante la autoridad competente. Canal de contacto: soporte@happypaws.co o el correo oficial que defina el responsable. El registro y uso de la plataforma implica aceptación de este tratamiento de datos conforme a la ley vigente. Plazo de conservación: mientras subsista la relación contractual y durante el tiempo exigido por la normatividad aplicable. El titular podrá ejercer sus derechos por los canales oficiales habilitados.')
add_note(doc, 'El aviso de privacidad debe mostrarse antes del registro y quedar almacenado como evidencia de aceptación.')
subsection_title(doc, '7.3', 'Ley 1774 de 2016 - Bienestar Animal')
add_paragraph(doc, 'La plataforma debe alinearse con la Ley 1774 de 2016, que reconoce a los animales como seres sintientes y refuerza el deber de protección contra el maltrato.')
add_table(doc, ['Obligación', 'Aplicación en Happy Paws'], [
    ['Reportar maltrato', 'Debe existir un flujo de escalamiento a autoridades competentes'],
    ['No promover conductas ilegales', 'Queda prohibido publicar peleas de animales, zoofilia o contenidos similares'],
    ['Trazabilidad', 'Todo reporte debe conservar fecha, evidencia, ubicación y seguimiento'],
    ['Bienestar animal', 'Los módulos deben promover adopción responsable, atención y cuidado']
])
subsection_title(doc, '7.4', 'Contrato de suscripción para veterinarias')
add_paragraph(doc, 'Contrato de Suscripción - cláusulas mínimas:')
add_table(doc, ['Cláusula', 'Contenido mínimo'], [
    ['1. Objeto', 'Uso de la plataforma Happy Paws para la gestión de pacientes, citas, adopciones y reportes'],
    ['2. Alcance del servicio', 'Módulos contratados, soporte incluido y límites de uso'],
    ['3. Obligaciones del cliente', 'Usar la plataforma conforme a la ley, custodiar credenciales y autorizar tratamiento de datos'],
    ['4. Obligaciones del proveedor', 'Disponibilidad razonable, soporte, seguridad y protección de datos'],
    ['5. Terminación', 'Causales de suspensión, cancelación, incumplimiento y tratamiento de información al cierre'],
    ['6. Confidencialidad', 'Protección de información del negocio y de los titulares'],
    ['7. Tratamiento de datos', 'Autorización, finalidades y deberes legales'],
    ['8. Responsabilidad', 'Limitaciones razonables y canales de soporte']
])
subsection_title(doc, '7.5', 'Términos y condiciones generales')
add_paragraph(doc, 'La estructura mínima del sitio web debe incluir: objeto del servicio, aceptación de condiciones, registro y cuenta de usuario, uso permitido, restricciones, propiedad intelectual, confidencialidad, tratamiento de datos, limitación de responsabilidad, suspensión del servicio, reclamaciones y ley aplicable.')
subsection_title(doc, '7.6', 'Registro ante la SIC')
add_paragraph(doc, 'Para efectos de este proyecto, el registro ante la SIC y el cumplimiento del Registro Nacional de Bases de Datos debe considerarse obligatorio cuando la base supere 5.000 titulares o cuando la operación crezca a un volumen similar que exija formalización adicional.')
add_table(doc, ['Aspecto', 'Detalle'], [
    ['Autoridad', 'Superintendencia de Industria y Comercio (SIC)'],
    ['Obligación', 'Inscripción y mantenimiento del registro de bases de datos cuando aplique'],
    ['Sanción por incumplimiento', 'Hasta 2.000 salarios mínimos legales mensuales vigentes, según la ley aplicable']
])

# Seguridad
section_title(doc, 8, 'Seguridad adicional y cumplimiento técnico')
add_table(doc, ['Medida', 'Implementación'], [
    ['Rate limiting', 'Máximo 100 peticiones por minuto por IP'],
    ['CORS', 'Solo dominios autorizados y ambientes controlados'],
    ['Validación de entradas', 'Validar en backend y frontend, no solo en UI'],
    ['Logs de auditoría', 'Registrar quién hizo qué, cuándo y desde dónde'],
    ['Backups automáticos', 'Copias diarias con verificación de restauración'],
    ['Contraseñas', 'Hash con bcrypt o equivalente seguro']
])
add_table(doc, ['Riesgo técnico', 'Controles mínimos'], [
    ['Inyección SQL', 'Consultas parametrizadas, ORM, validación y pruebas básicas'],
    ['XSS', 'Escape de salida, sanitización y Content Security Policy'],
    ['Pérdida de datos', 'Backups, monitoreo y plan de restauración'],
    ['Acceso no autorizado', 'JWT, expiración, roles y protección de rutas']
])

# Story points
section_title(doc, 9, 'Story points realistas')
add_paragraph(doc, 'La estimación anterior de 122 SP en 8 semanas era demasiado ambiciosa para un equipo junior-mid. Se ajusta a una distribución realista de 108 SP y a una velocidad coherente con un equipo pequeño que trabaja en iteraciones cortas.')
add_table(doc, ['Sprint', 'SP asignados', 'Justificación'], [
    ['Sprint 1', '25 SP', 'Base, autenticación y estructura técnica'],
    ['Sprint 2', '28 SP', 'Mascotas, agenda y trazabilidad'],
    ['Sprint 3', '30 SP', 'Adopciones, reportes y flujo social'],
    ['Sprint 4', '25 SP', 'Panel administrativo, seguridad y despliegue'],
    ['Total', '108 SP', 'Velocidad realista para el plazo y el tamaño del equipo']
])

# Endpoints
section_title(doc, 10, 'Endpoints REST y códigos de error')
add_table(doc, ['Método', 'Endpoint', 'Descripción'], [
    ['POST', '/api/v1/auth/register', 'Registro de usuario'],
    ['POST', '/api/v1/auth/login', 'Inicio de sesión'],
    ['POST', '/api/v1/auth/refresh', 'Refrescar token JWT'],
    ['GET', '/api/v1/pets', 'Listar mascotas'],
    ['POST', '/api/v1/pets', 'Crear mascota'],
    ['GET', '/api/v1/pets/{id}', 'Detalle de mascota'],
    ['PUT', '/api/v1/pets/{id}', 'Actualizar mascota'],
    ['DELETE', '/api/v1/pets/{id}', 'Eliminar mascota'],
    ['GET', '/api/v1/pets/{id}/history', 'Historial clínico'],
    ['GET', '/api/v1/vets', 'Listar veterinarias'],
    ['GET', '/api/v1/appointments', 'Listar citas'],
    ['POST', '/api/v1/appointments', 'Crear cita'],
    ['PATCH', '/api/v1/appointments/{id}/reschedule', 'Reprogramar cita'],
    ['PATCH', '/api/v1/appointments/{id}/cancel', 'Cancelar cita'],
    ['GET', '/api/v1/adoptions', 'Listar adopciones'],
    ['POST', '/api/v1/adoptions', 'Publicar adopción'],
    ['POST', '/api/v1/adoptions/{id}/apply', 'Postular a adopción'],
    ['PATCH', '/api/v1/adoptions/{id}/approve', 'Aprobar postulación'],
    ['PATCH', '/api/v1/adoptions/{id}/reject', 'Rechazar postulación'],
    ['GET', '/api/v1/reports', 'Listar reportes'],
    ['POST', '/api/v1/reports', 'Crear reporte'],
    ['GET', '/api/v1/admin/stats', 'Métricas admin'],
    ['GET', '/api/v1/admin/users', 'Gestión de usuarios']
])
add_table(doc, ['Código', 'Significado', 'Ejemplo'], [
    ['200 OK', 'Petición exitosa', 'GET /api/v1/pets'],
    ['201 Created', 'Recurso creado', 'POST /api/v1/pets'],
    ['400 Bad Request', 'Datos inválidos', 'Email mal formado'],
    ['401 Unauthorized', 'No autenticado', 'Token faltante'],
    ['403 Forbidden', 'Sin permisos', 'Usuario común accede a /admin'],
    ['404 Not Found', 'Recurso no existe', 'Pet ID inválido'],
    ['500 Internal Error', 'Error del servidor', 'Excepción no manejada']
])

# Competencia
section_title(doc, 11, 'Análisis de competencia')
add_paragraph(doc, 'Se agrega una comparación breve con software veterinario conocido en Colombia y la propuesta Happy Paws, para mostrar el espacio competitivo y el diferencial del proyecto.')
add_table(doc, ['Competidor', 'Precio', 'Ventajas', 'Desventajas'], [
    ['Vetmanager', 'COP 180.000 - 350.000/mes', 'Establecido, muchas funciones', 'Caro, no tiene adopciones ni reportes ciudadanos'],
    ['G247', 'COP 150.000 - 280.000/mes', 'Buen soporte', 'Curva de aprendizaje alta'],
    ['MVZ Cloud', 'COP 70.000 - 150.000/mes', 'Colombiano, económico', 'Poca escalabilidad'],
    ['VetSys', 'COP 100.000 - 200.000/mes', 'Opción conocida en el mercado', 'Menor foco en bienestar animal'],
    ['Happy Paws', 'COP 79.000 - 299.000/mes', 'Adopciones + reportes + bienestar animal', 'Nuevo en el mercado']
])

# Riesgos
section_title(doc, 12, 'Análisis de riesgos legales y técnicos')
add_table(doc, ['Riesgo', 'Probabilidad', 'Impacto', 'Mitigación'], [
    ['Incumplimiento de protección de datos (Ley 1581)', 'Media', 'Alto', 'Incluir aviso de privacidad, obtener consentimiento y firmar contrato de encargo'],
    ['Denuncias por maltrato animal no gestionadas', 'Media', 'Alto', 'Establecer flujo de escalamiento a autoridades'],
    ['Fuga de información de mascotas y dueños', 'Baja', 'Crítico', 'Encriptación, JWT y backups cifrados'],
    ['Competencia con software establecido', 'Alta', 'Medio', 'Enfoque en adopciones y reportes diferenciales'],
    ['Retraso en aprobaciones del cliente', 'Alta', 'Medio', 'Definir validaciones por sprint y responsables de aceptación'],
    ['Rotación del equipo técnico', 'Media', 'Alto', 'Documentación viva, pares de respaldo y backlog priorizado'],
    ['Dependencia de APIs de terceros', 'Media', 'Medio', 'Fallbacks, contratos de servicio y monitoreo de integraciones'],
    ['Escalabilidad inesperada', 'Baja', 'Alto', 'Pruebas de carga, colas y capacidad incremental']
])

# DoD
section_title(doc, 13, 'Definición de done mejorada')
add_table(doc, ['Criterio', 'Descripción'], [
    ['Código revisado', 'Aprobado por al menos 2 desarrolladores'],
    ['Pruebas funcionales', 'Flujos principales probados sin fallas críticas'],
    ['Pruebas de seguridad', 'Validación básica contra SQL Injection y XSS'],
    ['Documentación', 'API actualizada y changelog mínimo'],
    ['Calidad', 'Sin errores críticos de severidad 1 o 2'],
    ['Aprobación del PO', 'Aceptación formal del entregable']
])
add_note(doc, 'No se considera terminado un Sprint si existe un bloqueo de seguridad, documentación incompleta o deuda técnica crítica sin plan de cierre.')

# No MVP
section_title(doc, 14, 'Qué no incluir en el MVP')
add_paragraph(doc, 'Mantener el alcance del MVP es tan importante como construirlo. Se conserva la lista de exclusiones para evitar sobrecosto y desvío funcional.')
add_table(doc, ['Funcionalidad excluida', 'Razón'], [
    ['Telemedicina', 'Requiere infraestructura costosa y validación clínica adicional'],
    ['E-commerce', 'Añade complejidad logística innecesaria'],
    ['Microservicios', 'El MVP debe iniciar como monolito modular'],
    ['Gamificación', 'No es prioritaria para validar el mercado inicial']
])

# Conclusión
section_title(doc, 15, 'Conclusión y nota final')
add_paragraph(doc, 'Happy Paws queda consolidado como una propuesta viable para digitalizar el sector veterinario con enfoque social, técnico, comercial y legal. El documento corregido ya incluye costos realistas, cronograma ajustado, legalidad colombiana, seguridad base, competidores y una definición de alcance ejecutable.')
add_table(doc, ['Siguiente paso', 'Responsable', 'Resultado esperado'], [
    ['Validar alcance final', 'Cliente + equipo', 'Backlog cerrado y priorizado'],
    ['Aprobar presupuesto', 'Cliente', 'Contrato y presupuesto listos'],
    ['Iniciar desarrollo', 'Equipo TI', 'MVP en construcción'],
    ['Abrir piloto', 'Cliente + usuarios', 'Validación con usuarios reales']
])
add_paragraph(doc, 'Documento Happy Paws Versión 3.1 - Definitiva Corregida - Aprobado para uso ejecutivo, legal y comercial.')

# Anexos
section_title(doc, 16, 'Anexos')
add_table(doc, ['Anexo', 'Contenido'], [
    ['A1', 'Supuestos financieros y fórmula de costos'],
    ['A2', 'Aviso de privacidad completo'],
    ['A3', 'Contrato de suscripción completo'],
    ['A4', 'Backlog resumido y criterios de aceptación'],
    ['A5', 'Mapa de arquitectura técnica']
])

subsection_title(doc, 'A2', 'Aviso de privacidad')
add_paragraph(doc, 'AVISO DE PRIVACIDAD. Responsable del tratamiento: Happy Paws SAS o la persona jurídica que opere la plataforma para la veterinaria aliada. Finalidades: gestionar usuarios, mascotas, citas, adopciones, reportes, comunicaciones operativas, soporte, seguridad, analítica y cumplimiento legal. Datos recolectados: nombre, cédula o documento, correo, teléfono, dirección, datos de mascotas, historial clínico y registros asociados al uso del sistema. Derechos: consultar, actualizar, rectificar, suprimir, revocar la autorización y presentar quejas ante la autoridad competente. Canal de contacto: soporte@happypaws.co o el correo oficial que defina el responsable. El registro y uso de la plataforma implica aceptación de este tratamiento de datos conforme a la ley vigente. Plazo de conservación: mientras subsista la relación contractual y durante el tiempo exigido por la normatividad aplicable. El titular podrá ejercer sus derechos por los canales oficiales habilitados.')

subsection_title(doc, 'A3', 'Contrato de suscripción')
add_paragraph(doc, 'CONTRATO DE SUSCRIPCIÓN. Cláusula 1. Objeto: el presente contrato regula el uso de la plataforma Happy Paws para la gestión de pacientes, citas, adopciones y reportes. Cláusula 2. Alcance del servicio: el proveedor habilita los módulos contratados, soporte funcional básico, actualizaciones menores y acompañamiento de puesta en marcha. Cláusula 3. Obligaciones del cliente: custodiar credenciales, garantizar el uso lícito del sistema y obtener las autorizaciones necesarias para tratar datos personales. Cláusula 4. Obligaciones del proveedor: mantener medidas razonables de disponibilidad, seguridad, respaldo y atención de incidentes. Cláusula 5. Terminación: cualquiera de las partes podrá terminar el contrato por incumplimiento, vencimiento o decisión mutua, preservando la trazabilidad y la información exigida por ley. Cláusula 6. Confidencialidad: las partes se obligan a no divulgar información técnica, comercial o personal a la que accedan por razón del servicio. Cláusula 7. Tratamiento de datos: el cliente autoriza el tratamiento de datos personales conforme a la Ley 1581 de 2012 y normas relacionadas. Cláusula 8. Responsabilidad: el proveedor no responderá por usos indebidos del sistema fuera del alcance contratado. Cláusula 9. Soporte: los canales y tiempos de respuesta se definirán en la orden de servicio o anexo comercial. Cláusula 10. Ley aplicable: este contrato se rige por la legislación colombiana.')

add_paragraph(doc, 'A1. Supuestos financieros: costo base de recursos humanos = ' + format_cop(BASE_COST) + '; contingencia = ' + format_cop(CONTINGENCY) + '; costo total = ' + format_cop(TOTAL_COST) + '; precio sugerido = ' + format_cop(SALE_PRICE) + '.')

add_paragraph(doc, 'A4. Backlog resumido: autenticación, mascotas, citas, adopciones, reportes, panel administrativo, auditoría, seguridad y despliegue.')
add_paragraph(doc, 'A5. Mapa de arquitectura técnica: frontend en React, backend en Spring Boot, base relacional en PostgreSQL, autenticación con JWT, despliegue en Docker y almacenamiento de imágenes en Cloudinary.')

# Cierre final
add_paragraph(doc, 'La siguiente validación recomendada es revisar el archivo generado en Word para confirmar formato final, tabla de contenido e incorporación visual de los anexos.')

doc.save(OUTPUT_FILE)
print(f'Generado: {OUTPUT_FILE}')
